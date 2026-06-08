"""สร้าง abstract.docx สำหรับ BoxBox"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Font helpers ─────────────────────────────────────────────────────────────
FONT = 'TH Sarabun New'
FONT_EN = 'TH Sarabun New'

def set_run(run, bold=False, size=16, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.name = FONT
    run.font.size = Pt(size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),      FONT)
    rFonts.set(qn('w:hAnsi'),      FONT)
    rFonts.set(qn('w:cs'),         FONT)
    rFonts.set(qn('w:eastAsia'),   FONT)
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text='', bold=False, size=16, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=4, color=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(22)
    if text:
        run = p.add_run(text)
        set_run(run, bold=bold, size=size, color=color, italic=italic)
    return p

def add_mixed(parts, size=16, space_before=0, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
    """parts = list of (text, bold, color)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(22)
    for text, bold, color in parts:
        run = p.add_run(text)
        set_run(run, bold=bold, size=size, color=color)
    return p

def add_heading(text, level=1):
    if level == 1:
        p = add_para(text, bold=True, size=18, space_before=10, space_after=4,
                     color=(31, 41, 123))
        # underline
        for run in p.runs:
            run.underline = True
    elif level == 2:
        p = add_para(text, bold=True, size=16, space_before=8, space_after=2,
                     color=(63, 63, 191))
    else:
        p = add_para(text, bold=True, size=16, space_before=4, space_after=2,
                     color=(80, 80, 80))
    return p

def add_note(text):
    """ข้อความสีส้ม = ส่วนที่ต้องเพิ่ม"""
    p = add_para()
    run = p.add_run('⚠️  ' + text)
    set_run(run, bold=True, size=15, color=(180, 80, 0), italic=True)
    p.paragraph_format.left_indent = Cm(1)
    return p

def add_bullet(text, bold_prefix=None, size=16, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(22)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ': ')
        set_run(r1, bold=True, size=size)
    r2 = p.add_run(text)
    set_run(r2, bold=False, size=size)
    return p

def add_bullet_note(text, size=15):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(22)
    r = p.add_run('⚠️  ' + text)
    set_run(r, bold=True, size=size, color=(180, 80, 0), italic=True)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
add_para('บทคัดย่อ', bold=True, size=22,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)

add_para('BoxBox: ระบบบริหารจัดการกล่องยาฉุกเฉินประจำหอผู้ป่วยแบบดิจิทัล',
         bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=0, space_after=2, color=(31, 41, 123))

add_para('Ward Emergency Drug Box Management System',
         bold=False, size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
         italic=True, space_before=0, space_after=2, color=(80,80,80))

add_note('ระบุชื่อโรงพยาบาล / หน่วยงาน / ชื่อผู้พัฒนา / ปีที่พัฒนา ที่บรรทัดนี้')

add_para()

# ══════════════════════════════════════════════════════════════════════════════
# 1. บทนำ
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1.  บทนำ  (หลักการและเหตุผล  แนวคิด  วัตถุประสงค์)', level=1)

add_heading('1.1  หลักการและเหตุผล', level=2)

add_para(
    'กล่องยาฉุกเฉินประจำหอผู้ป่วย (Ward Emergency Drug Box) เป็นองค์ประกอบสำคัญของระบบยา'
    'ในโรงพยาบาล ซึ่งต้องมีความพร้อมใช้ตลอด 24 ชั่วโมง กล่องยาฉุกเฉินครอบคลุมหลายประเภท ได้แก่ '
    'กล่อง CPR ผู้ใหญ่ กล่อง CPR ทารกแรกเกิด กล่อง PPH กล่อง PIH กล่อง ACS กล่อง Emergency OR '
    'กล่อง Emergency ER และกล่อง EMS/Refer รวม 8 ประเภท ครอบคลุมหอผู้ป่วย 22 แห่ง '
    'ในโรงพยาบาล [⚠️ ระบุชื่อโรงพยาบาล]',
    size=16, space_after=4)

add_para(
    'อย่างไรก็ตาม กระบวนการบริหารจัดการที่ผ่านมายังคงพึ่งพาการบันทึกด้วยกระดาษและการตรวจสอบ'
    'ด้วยสายตาเป็นหลัก ซึ่งก่อให้เกิดปัญหาสำคัญหลายประการ ดังนี้',
    size=16, space_after=2)

issues = [
    ('ความเสี่ยงด้านคุณภาพยา', 'ยาในกล่องฉุกเฉินอาจหมดอายุโดยที่เจ้าหน้าที่ไม่ทราบ เนื่องจากขาดระบบแจ้งเตือนอัตโนมัติ'),
    ('ขาดการตรวจสอบย้อนหลัง', 'ไม่สามารถติดตาม Lot Number ของยา ผู้บรรจุ และประวัติการใช้งานกล่องแต่ละกล่องได้อย่างเป็นระบบ'),
    ('กระบวนการพิมพ์ฉลากและปกกล่องไม่เป็นมาตรฐาน', 'รูปแบบการพิมพ์แตกต่างกันในแต่ละหอผู้ป่วย ทำให้การตรวจสอบทำได้ยาก'),
    ('ขาดข้อมูลสำหรับการบริหาร', 'ผู้บริหารและเภสัชกรไม่มีข้อมูล dashboard แบบ real-time สำหรับตัดสินใจ'),
    ('ยืนยันความพร้อมใช้ล่าช้า', 'พยาบาลต้องรายงานกลับเป็นกระดาษ ทำให้เภสัชกรไม่ทราบสถานะกล่องที่ส่งออกไป'),
]
for bold_part, text in issues:
    add_bullet(text, bold_prefix=bold_part)

add_note('เพิ่มข้อมูลเชิงประจักษ์: จำนวนครั้งที่พบยาหมดอายุในกล่องฉุกเฉินก่อนพัฒนาระบบ (ย้อนหลัง 1-2 ปี) '
         'เพื่อแสดงขนาดของปัญหา เช่น "ปี 2566 พบยาหมดอายุ X ครั้ง มูลค่า X บาท"')

add_heading('1.2  แนวคิดของนวัตกรรม', level=2)

add_para(
    'BoxBox ถูกพัฒนาขึ้นภายใต้แนวคิด "Digital Transformation of Ward Drug Safety" '
    'โดยมุ่งเน้นการแปลงกระบวนการจัดการกล่องยาฉุกเฉินทั้งหมดให้เป็นดิจิทัล ครอบคลุมตั้งแต่ '
    'การบรรจุยา → การพิมพ์เอกสาร → การจ่ายกล่องไปหอผู้ป่วย → การยืนยันความพร้อมใช้ → '
    'การแจ้งเตือนหมดอายุ → การรับคืนและบรรจุใหม่ → การรายงานผล '
    'โดยใช้เทคโนโลยีที่เข้าถึงได้และไม่มีค่าใช้จ่าย (Zero-cost Technology Stack)',
    size=16, space_after=4)

concepts = [
    'QR Code บน "ปกกล่อง" ให้พยาบาลสแกนยืนยันความพร้อมใช้ได้ทันทีผ่านสมาร์ทโฟน',
    'Real-time Dashboard พร้อม Color-coded Alert (แดง/เหลือง/เขียว) สำหรับติดตามสถานะกล่องทุกกล่องพร้อมกัน',
    'Cloud Sync ผ่าน Google Apps Script เพื่อสำรองข้อมูลและรับข้อมูล QR confirmation จากหอผู้ป่วย',
    'LINE Notify Alert อัตโนมัติเมื่อยาหรือกล่องใกล้หมดอายุตามเกณฑ์ที่กำหนด',
    'Print Template Engine แบบ drag-and-drop ที่ปรับแต่งได้ตามบริบทของแต่ละโรงพยาบาล',
    'รายงาน 7 รูปแบบ พร้อม Export xlsx สำหรับผู้บริหารและเภสัชกร',
]
for c in concepts:
    add_bullet(c)

add_heading('1.3  วัตถุประสงค์', level=2)

objectives = [
    'เพื่อพัฒนาระบบบริหารจัดการกล่องยาฉุกเฉินประจำหอผู้ป่วยแบบดิจิทัลที่ครอบคลุมทุกกระบวนการในวงจรชีวิตของกล่องยา',
    'เพื่อลดความเสี่ยงจากการใช้ยาที่หมดอายุในสถานการณ์ฉุกเฉิน ด้วยระบบแจ้งเตือนอัตโนมัติแบบหลายช่องทาง',
    'เพื่อเพิ่มประสิทธิภาพและความโปร่งใสในกระบวนการตรวจสอบ ด้วยการติดตาม Lot Number และประวัติผู้บรรจุทุก cycle',
    'เพื่อยืนยันความพร้อมใช้ของกล่องยาที่หอผู้ป่วยได้สะดวกและรวดเร็วผ่าน QR Code โดยไม่ต้องรายงานกระดาษ',
    'เพื่อพัฒนาแพลตฟอร์มที่สามารถปรับแต่ง ต่อยอด และขยายผลใช้งานในสถานพยาบาลระดับต่างๆ ได้โดยไม่มีค่าใช้จ่าย',
]
for i, obj in enumerate(objectives, 1):
    add_bullet(f'{obj}', bold_prefix=f'วัตถุประสงค์ที่ {i}')

add_para()

# ══════════════════════════════════════════════════════════════════════════════
# 2. วิธีการพัฒนา
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2.  วิธีการพัฒนานวัตกรรม/สิ่งประดิษฐ์', level=1)

add_heading('2.1  สถาปัตยกรรมของระบบ', level=2)

add_para(
    'BoxBox พัฒนาในรูปแบบ Hybrid Desktop Application โดยใช้แนวทาง Web-in-Native '
    'ซึ่งรวมข้อดีของ desktop application (ทำงานได้แบบ offline, เข้าถึง printer โดยตรง) '
    'เข้ากับข้อดีของ web technology (UI ทันสมัย, พัฒนารวดเร็ว, ไม่ต้องติดตั้ง runtime เพิ่ม)',
    size=16, space_after=4)

tech_stack = [
    ('Shell', 'C# WinForms + Microsoft WebView2 (.NET 8) — รองรับ Windows 10/11'),
    ('UI Framework', 'React 18 via CDN + Babel Standalone — ไม่ต้องมี build pipeline หรือ Node.js'),
    ('Local Storage', 'Browser localStorage — ทำงานได้แบบ offline สมบูรณ์'),
    ('Cloud Backend', 'Google Apps Script (GAS) — ฟรี, sync ข้อมูล, รับ QR confirmation'),
    ('Print Engine', 'GDI+ via WebBridge (C#) + HTML/CSS Print Template'),
    ('QR Code', 'api.qrserver.com (generation) + GAS webhook (confirmation)'),
    ('Notification', 'LINE Notify API + แจ้งเตือนใน Dashboard'),
    ('Export', 'SheetJS (xlsx) สำหรับ Export รายงาน'),
]
for bold_part, text in tech_stack:
    add_bullet(text, bold_prefix=bold_part, size=15)

add_heading('2.2  กระบวนการพัฒนา', level=2)

steps = [
    ('ขั้นตอนที่ 1: วิเคราะห์ความต้องการ',
     'จัดประชุมระบบยาร่วมกับเภสัชกร พยาบาล และผู้บริหาร (ประชุมระบบยา ครั้งที่ 1/2568) '
     'เพื่อกำหนด 8 ประเภทกล่องยา พร้อมรายการยาใน template แต่ละกล่อง '
     'และ 22 หอผู้ป่วยที่เกี่ยวข้อง'),
    ('ขั้นตอนที่ 2: ออกแบบ Data Model',
     'กำหนด lifecycle ของกล่องยา: สร้าง → filling → ready → dispatched → retiring '
     'และโครงสร้างข้อมูล localStorage 8 keys'),
    ('ขั้นตอนที่ 3: พัฒนา Core UI',
     'พัฒนา React components: Dashboard, FillModal (3 ขั้นตอน), Exchange, Settings, Report '
     'โดยใช้ Babel Standalone compileแบบ runtime ไม่ต้องมี build step'),
    ('ขั้นตอนที่ 4: พัฒนา Print System',
     'สร้าง PrintTemplates engine รองรับ 3 รูปแบบ (ใบรายการยา, สติกเกอร์, ปกกล่อง) '
     'พร้อม Canvas Editor drag-and-drop สำหรับปรับ layout'),
    ('ขั้นตอนที่ 5: พัฒนา Cloud Integration',
     'เชื่อมต่อ GAS สำหรับ sync ข้อมูล, รับ QR confirmation, และส่ง LINE Notify alert'),
    ('ขั้นตอนที่ 6: ทดสอบและปรับปรุง',
     '[⚠️ ระบุวิธีทดสอบ เช่น User Acceptance Test กับเภสัชกร X คน พยาบาล X คน '
     'ระยะเวลา X เดือน จำนวนกล่องที่ทดสอบ X กล่อง]'),
]
for bold_part, text in steps:
    if '⚠️' in text:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = Pt(22)
        r1 = p.add_run(bold_part + ': ')
        set_run(r1, bold=True, size=16)
        r2 = p.add_run('[⚠️ ระบุวิธีทดสอบ เช่น User Acceptance Test กับเภสัชกร X คน พยาบาล X คน '
                       'ระยะเวลา X เดือน จำนวนกล่องที่ทดสอบ X กล่อง]')
        set_run(r2, bold=True, size=15, color=(180, 80, 0), italic=True)
    else:
        add_bullet(text, bold_prefix=bold_part)

add_heading('2.3  ฟีเจอร์และความสามารถของระบบ', level=2)

features = [
    ('Dashboard แบบ Real-time',
     'แสดงสถานะกล่องทั้งหมดด้วย Color-coded Badge (🔴 วิกฤต / 🟡 เฝ้าระวัง / 🟢 ปกติ) '
     'พร้อม filter ตามหมวดหมู่ ประเภทกล่อง และสถานะ'),
    ('FillModal 3 ขั้นตอน',
     'บรรจุยา (pre-fill จากครั้งก่อน) → เลือกกล่องแทนที่ → พิมพ์ปก+ยืนยันโดยเภสัชกร '
     'บันทึก Lot Number ทุกรายการยา'),
    ('QR Confirmation System',
     'ปกกล่องมี QR Code เชื่อมกับ GAS webhook พยาบาลสแกนผ่านสมาร์ทโฟนเพื่อยืนยันความพร้อมใช้ '
     'ข้อมูลซิงค์กลับมายังโปรแกรมอัตโนมัติ'),
    ('Print Template Engine',
     '3 รูปแบบ: ใบรายการยา (A4), สติกเกอร์ (ปรับขนาดได้), ปกกล่อง (แนวนอน 1/5 A4) '
     'พร้อม Canvas Editor drag-and-drop — ใช้ฟอนต์ Sarabun รองรับภาษาไทยสมบูรณ์'),
    ('ระบบรายงาน 7 รูปแบบ',
     'บรรจุ+จ่าย / หมดอายุ / แจ้งเตือน / ใกล้หมดอายุ (กล่อง+ยา) / ยาที่ใช้ / ความพร้อมใช้ / ประวัติกล่อง '
     'ทุกรายงาน export xlsx ได้'),
    ('Multi-channel Alert',
     'แจ้งเตือนผ่าน Dashboard badge, LINE Notify (กำหนด target ได้หลายกลุ่ม), '
     'และตรวจสอบอัตโนมัติเมื่อเปิดโปรแกรม'),
    ('Lot Number Tracking',
     'บันทึกและแสดง Lot No. ของยาทุกรายการในใบรายการยา สติกเกอร์ รายงาน และประวัติกล่อง '
     'รองรับการ recall ยาตาม Lot No.'),
    ('Cloud Sync (GAS)',
     'sync ข้อมูลหลัก 13 keys ไป Google Sheets, รับ QR confirmation, '
     'สำรองข้อมูลอัตโนมัติ ทำงานได้ทั้ง online/offline'),
]
for bold_part, text in features:
    add_bullet(text, bold_prefix=bold_part)

add_heading('2.4  ความแตกต่างจากระบบอื่นในประเภทเดียวกัน', level=2)

diff_table = doc.add_table(rows=1, cols=4)
diff_table.style = 'Table Grid'
hdr = diff_table.rows[0].cells
labels = ['คุณสมบัติ', 'BoxBox', 'ระบบ Manual (กระดาษ)', 'ระบบ HIS ทั่วไป']
for i, (cell, label) in enumerate(zip(hdr, labels)):
    p = cell.paragraphs[0]
    run = p.add_run(label)
    set_run(run, bold=True, size=14, color=(255,255,255) if i==1 else (255,255,255))
    from docx.oxml.ns import qn as _qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(_qn('w:fill'), '1F297B' if i==0 else ('2563EB' if i==1 else ('6B7280' if i==2 else '374151')))
    shd.set(_qn('w:color'), 'auto')
    shd.set(_qn('w:val'), 'clear')
    tcPr.append(shd)

rows_data = [
    ('QR Confirmation', '✅ สแกนยืนยันพร้อมใช้', '❌ ไม่มี', '❌ ไม่มีสำหรับกล่องฉุกเฉิน'),
    ('แจ้งเตือน LINE', '✅ อัตโนมัติ', '❌ ไม่มี', '⚠️ บางระบบมี (ราคาสูง)'),
    ('Print Template', '✅ ปรับแต่งได้', '❌ ไม่มี', '⚠️ มีบ้างแต่ปรับแต่งไม่ได้'),
    ('Lot No. Tracking', '✅ ทุกรายการ', '⚠️ บางส่วน', '⚠️ ขึ้นกับ configuration'),
    ('Offline Mode', '✅ ทำงานได้เต็ม', '✅ ทำงานได้', '❌ ต้องต่อเครือข่ายโรงพยาบาล'),
    ('ค่าใช้จ่าย', '✅ ฟรี (Open Source)', '✅ ต่ำ แต่เสียเวลา', '❌ สูง (License ปีละแสน+)'),
    ('Export xlsx', '✅ ทุกรายงาน', '❌ ไม่มี', '⚠️ บางระบบมี'),
]
for row_data in rows_data:
    row = diff_table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, row_data)):
        p = cell.paragraphs[0]
        run = p.add_run(val)
        set_run(run, bold=(i==0), size=14)
        if i == 1:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'EFF6FF')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)

add_para()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ประโยชน์และการนำไปใช้
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3.  ประโยชน์และการนำไปใช้', level=1)

add_heading('3.1  ประโยชน์เชิงความปลอดภัยและคุณภาพ', level=2)

safety = [
    'ลดความเสี่ยงจากยาหมดอายุในกล่องฉุกเฉิน ด้วยระบบแจ้งเตือนล่วงหน้า 2 ระดับ (เฝ้าระวัง/วิกฤต)',
    'สนับสนุนมาตรฐาน HA (Hospital Accreditation) และ JCI ด้านความปลอดภัยในการใช้ยาฉุกเฉิน (Medication Safety)',
    'เพิ่มความโปร่งใสในกระบวนการ: บันทึกผู้บรรจุยา เภสัชกรผู้ตรวจสอบ Lot No. และ timestamp ทุก cycle',
    'ระบบ QR Confirmation ช่วยให้เภสัชกรทราบสถานะความพร้อมใช้ของกล่องที่ส่งออกทุกกล่อง',
    'ติดตาม Lot Number เพื่อรองรับการ recall ยาได้อย่างรวดเร็วและแม่นยำ',
]
for s in safety:
    add_bullet(s)

add_heading('3.2  ประโยชน์เชิงประสิทธิภาพและเศรษฐศาสตร์', level=2)

econ = [
    'ลดเวลาในการตรวจสอบและบันทึกข้อมูลด้วยมือ',
    'ลดมูลค่ายาที่สูญเสียจากการหมดอายุโดยไม่ทราบ',
    'Zero Licensing Cost: ใช้ .NET (ฟรี), React (ฟรี), Google Apps Script (ฟรี) ทำให้โรงพยาบาลไม่มีค่าใช้จ่ายซอฟต์แวร์',
    'ติดตั้งง่ายบน Windows 10/11 มาตรฐาน ไม่ต้องการ server หรือโครงสร้างพื้นฐานพิเศษ',
]
for e in econ:
    add_bullet(e)
for hint in [
    'ระบุเวลาเฉลี่ยที่ประหยัดได้ต่อเดือนสำหรับเภสัชกร (เช่น "ลดเวลา X ชั่วโมง/เดือน")',
    'ระบุมูลค่ายาที่ลดการสูญเสียได้ (เช่น "ลดมูลค่ายาหมดอายุจาก X บาท/ปี เหลือ Y บาท/ปี")',
    'ระบุต้นทุนรวมของการพัฒนา (ค่าแรงโปรแกรมเมอร์) เพื่อเปรียบเทียบกับระบบเชิงพาณิชย์',
]:
    add_bullet_note(hint)

add_heading('3.3  การนำไปใช้ในปัจจุบันและแผนต่อยอด', level=2)

current = [
    'ปัจจุบันใช้งานใน [⚠️ ระบุชื่อโรงพยาบาล] ครอบคลุม [⚠️ X] หอผู้ป่วย [⚠️ X] กล่องยาฉุกเฉิน',
    'บุคลากรที่ใช้งาน: เภสัชกร [⚠️ X] คน พยาบาล [⚠️ X] คน (สแกน QR)',
]
for c in current:
    if '⚠️' in c:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.line_spacing = Pt(22)
        r = p.add_run(c)
        set_run(r, bold=True, size=15, color=(180, 80, 0), italic=True)
    else:
        add_bullet(c)

future_plans = [
    ('ระยะสั้น', 'เพิ่มการนำเข้าข้อมูลยาจาก HIS โรงพยาบาลอัตโนมัติ (ลดการ key ซ้ำ)'),
    ('ระยะกลาง', 'พัฒนา Mobile App สำหรับพยาบาลในการยืนยันและตรวจสอบกล่องยา'),
    ('ระยะยาว', 'ขยายผลสู่โรงพยาบาลในเครือข่ายเขตสุขภาพ พร้อม multi-tenant cloud platform'),
    ('ต่อยอดงานวิจัย', 'ศึกษาผลลัพธ์ด้านความปลอดภัยของผู้ป่วย (Patient Safety Outcome) ก่อน-หลังใช้ระบบ'),
]
for bold_part, text in future_plans:
    add_bullet(text, bold_prefix=bold_part)

add_para()

# ══════════════════════════════════════════════════════════════════════════════
# 4. ภาพถ่าย / แบบจำลอง / การทดสอบ
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4.  ภาพถ่ายชิ้นงาน / แบบจำลอง / การทดสอบประสิทธิภาพ', level=1)

add_heading('4.1  แผนผังสถาปัตยกรรมระบบ (System Architecture)', level=2)

add_para(
    'BoxBox ใช้สถาปัตยกรรม Hybrid Desktop Application ดังแสดงในแผนผังด้านล่าง:',
    size=16, space_after=4)

arch_lines = [
    '┌─────────────────────────────────────────────────────────┐',
    '│               BoxBox Desktop Application                │',
    '│  ┌─────────────────────────────────────────────────┐   │',
    '│  │      C# WinForms Shell (.NET 8)                 │   │',
    '│  │  ┌───────────────────────────────────────────┐  │   │',
    '│  │  │    Microsoft WebView2 (Chromium Engine)   │  │   │',
    '│  │  │  ┌─────────────────────────────────────┐  │  │   │',
    '│  │  │  │   React 18 UI (CDN + Babel)         │  │  │   │',
    '│  │  │  │   Dashboard | Fill | Report | Print │  │  │   │',
    '│  │  │  └─────────────────────────────────────┘  │  │   │',
    '│  │  └───────────────────────────────────────────┘  │   │',
    '│  │  WebBridge (COM) ← Print / Alert / Version      │   │',
    '│  └─────────────────────────────────────────────────┘   │',
    '│               localStorage (Offline Storage)            │',
    '└──────────────────────┬──────────────────────────────────┘',
    '                       │ HTTPS (GAS Webhook)',
    '┌──────────────────────▼──────────────────────────────────┐',
    '│          Google Apps Script (Cloud Backend)             │',
    '│    Sync Data │ QR Confirmation │ LINE Notify            │',
    '└─────────────────────────────────────────────────────────┘',
]
p = doc.add_paragraph()
p.paragraph_format.line_spacing = Pt(16)
p.paragraph_format.space_after  = Pt(4)
for line in arch_lines:
    run = p.add_run(line + '\n')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

add_heading('4.2  ภาพหน้าจอระบบ (Screenshots)', level=2)

screenshots = [
    ('ภาพที่ 1', 'Dashboard Tab — แสดงกล่องยาทั้งหมดพร้อม Color-coded Badge และ KPI'),
    ('ภาพที่ 2', 'FillModal — ขั้นตอนที่ 1: บรรจุยาพร้อม Lot Number'),
    ('ภาพที่ 3', 'FillModal — ขั้นตอนที่ 3: Preview ปกกล่องพร้อม QR Code'),
    ('ภาพที่ 4', 'Print Template — ปกกล่อง (Cover Sheet) แนวนอน'),
    ('ภาพที่ 5', 'Print Template — ใบรายการยา (A4) พร้อมตาราง Lot No.'),
    ('ภาพที่ 6', 'Report Tab — รายงานความพร้อมใช้ (QR Confirmation)'),
    ('ภาพที่ 7', 'Template Designer — Canvas Editor drag-and-drop'),
    ('ภาพที่ 8', 'การสแกน QR Code ด้วยสมาร์ทโฟนที่หอผู้ป่วย'),
]
for bold_part, text in screenshots:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.line_spacing = Pt(22)
    r1 = p.add_run(bold_part + ': ')
    set_run(r1, bold=True, size=15)
    r2 = p.add_run(text + ' ')
    set_run(r2, size=15)
    r3 = p.add_run('[⚠️ แนบภาพจริงที่นี่]')
    set_run(r3, bold=True, size=14, color=(180, 80, 0), italic=True)

add_heading('4.3  การทดสอบประสิทธิภาพ', level=2)

add_note('ส่วนนี้สำคัญมากสำหรับคะแนน — ควรเพิ่มข้อมูลเชิงประจักษ์จากการทดสอบจริง:')

test_items = [
    'จำนวนกล่องยาที่ทดสอบ: [⚠️ X กล่อง] ครอบคลุม [⚠️ X ประเภท] ใน [⚠️ X หอผู้ป่วย]',
    'ระยะเวลาทดสอบ: [⚠️ X เดือน] ตั้งแต่ [⚠️ เดือน/ปี] ถึง [⚠️ เดือน/ปี]',
    'ความถูกต้องของการแจ้งเตือน: [⚠️ X%] (จำนวนครั้งที่แจ้งเตือนถูกต้อง / ทั้งหมด)',
    'Uptime ของระบบ: [⚠️ X%] (ระบบพร้อมใช้งาน / เวลาทั้งหมด)',
    'เวลาเฉลี่ยในการบรรจุยา 1 กล่อง ด้วยระบบ BoxBox: [⚠️ X นาที] เทียบกับระบบกระดาษ [⚠️ X นาที]',
    'ความพึงพอใจของผู้ใช้: [⚠️ X/5 คะแนน] จากการสำรวจเภสัชกร [⚠️ X] คน พยาบาล [⚠️ X] คน',
    'จำนวนครั้งที่ QR Confirmation ใช้งาน: [⚠️ X ครั้ง] ใน [⚠️ X วัน]',
]
for item in test_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.line_spacing = Pt(22)
    r = p.add_run(item)
    set_run(r, bold='⚠️' in item, size=15,
            color=(180, 80, 0) if '⚠️' in item else None,
            italic='⚠️' in item)

add_para()

# ══════════════════════════════════════════════════════════════════════════════
# 5. เกณฑ์คะแนน mapping
# ══════════════════════════════════════════════════════════════════════════════
add_heading('5.  สรุปการตอบสนองต่อเกณฑ์การพิจารณาผลงานนวัตกรรม', level=1)

criteria = [
    ('(1) แนวคิดเหมาะสมกับบริบท',
     'BoxBox ถูกพัฒนาตอบสนองความต้องการจริงจากการประชุมระบบยาของโรงพยาบาล '
     'ครอบคลุม 8 ประเภทกล่อง 22 หอผู้ป่วย ซึ่งสอดคล้องกับบริบทของโรงพยาบาลระดับ '
     '[⚠️ ระบุระดับ เช่น M2 / S / A]',
     False),
    ('(2) ความแปลกใหม่ โดดเด่น แตกต่าง',
     'นวัตกรรมแรกในระดับโรงพยาบาลที่รวม QR Confirmation + LINE Notify + '
     'Canvas Print Template + Cloud Sync ไว้ในระบบเดียว โดยไม่มีค่าใช้จ่าย '
     'และทำงาน offline ได้สมบูรณ์',
     False),
    ('(3) เพิ่มมูลค่า / ลดต้นทุน',
     '[⚠️ ระบุตัวเลขจริง] ลดมูลค่ายาหมดอายุ X บาท/ปี | ลดต้นทุนเมื่อเทียบกับซอฟต์แวร์เชิงพาณิชย์ที่มีราคา > 100,000 บาท/ปี',
     True),
    ('(4) ความชาญฉลาดทางเทคนิค',
     'ใช้ Hybrid Architecture (C# + WebView2 + React) ที่ช่วยให้พัฒนาเร็ว บำรุงรักษาง่าย '
     'และใช้ได้กับ Windows มาตรฐาน | ไม่ต้องมี server | ทำงาน offline ได้สมบูรณ์',
     False),
    ('(5) การออกแบบสวยงาม เหมาะสมกับการใช้งาน',
     'UI ใช้ React 18 + Tailwind-inspired CSS พร้อม Color-coded Badge | '
     'ฟอนต์ TH Sarabun New สำหรับเอกสารภาษาไทย | '
     'Canvas Editor drag-and-drop สำหรับ print template',
     False),
    ('(6) นำไปใช้ประโยชน์และต่อยอดได้',
     'Deploy ได้ทันทีบน Windows มาตรฐาน | '
     'Seed data ปรับตามบริบทโรงพยาบาลได้ | '
     'แผนต่อยอด: Mobile App, HIS Integration, Multi-hospital Platform',
     False),
]

crit_table = doc.add_table(rows=1, cols=3)
crit_table.style = 'Table Grid'
hdr2 = crit_table.rows[0].cells
for cell, label, fill_hex in zip(hdr2, ['เกณฑ์', 'การตอบสนองของ BoxBox', 'สถานะ'], ['1F297B','1F297B','1F297B']):
    p = cell.paragraphs[0]
    run = p.add_run(label)
    set_run(run, bold=True, size=14, color=(255,255,255))
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

for criterion, response, needs_data in criteria:
    row = crit_table.add_row()
    cells = row.cells
    p0 = cells[0].paragraphs[0]
    r0 = p0.add_run(criterion)
    set_run(r0, bold=True, size=13)

    p1 = cells[1].paragraphs[0]
    r1 = p1.add_run(response)
    set_run(r1, bold=needs_data, size=13,
            color=(180, 80, 0) if needs_data else None,
            italic=needs_data)

    p2 = cells[2].paragraphs[0]
    status = '⚠️ ต้องเพิ่มข้อมูล' if needs_data else '✅ พร้อม'
    r2 = p2.add_run(status)
    set_run(r2, bold=True, size=13,
            color=(180, 80, 0) if needs_data else (22, 163, 74))

add_para()

# ══════════════════════════════════════════════════════════════════════════════
# คำอธิบายสัญลักษณ์ที่ใช้
# ══════════════════════════════════════════════════════════════════════════════
add_heading('หมายเหตุ: สัญลักษณ์ที่ใช้ในเอกสารนี้', level=1)

notes = [
    '✅  ข้อมูลพร้อมใช้งาน ไม่ต้องแก้ไข',
    '⚠️  [ข้อความสีส้ม] = ส่วนที่ต้องเพิ่มข้อมูลจริงก่อนส่งบทคัดย่อ',
    'ข้อมูลที่สำคัญที่สุด: ตัวเลขผลลัพธ์จากการทดสอบจริง (ข้อ 4.3) '
    'และข้อมูลเชิงเศรษฐศาสตร์ (ข้อ 3.2) จะมีผลมากต่อคะแนนเกณฑ์ที่ (3)',
]
for note in notes:
    add_bullet(note, size=15)

add_para()
add_para('─' * 80, size=10, color=(150,150,150))
add_note('เอกสารนี้สร้างโดย Claude Code โดยอัตโนมัติจาก source code ของ BoxBox — '
         'กรุณาตรวจสอบและเพิ่มเติมข้อมูลในส่วนที่มีเครื่องหมาย ⚠️ ก่อนส่ง')

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
doc.save(r'D:\COACH\source\repo\BoxBox\abstract.docx')
print('abstract.docx created successfully')
